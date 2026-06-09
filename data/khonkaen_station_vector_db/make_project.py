from docx import Document
from pathlib import Path
import json, re, joblib
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.neighbors import NearestNeighbors

BASE=Path(__file__).parent
SRC=Path('/mnt/data/metadata_for_layout.docx')

def read_docx(path):
    doc=Document(path)
    parts=[]
    for para in doc.paragraphs:
        t=' '.join(para.text.split())
        if t: parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells=[' '.join(c.text.split()) for c in row.cells]
            if any(cells): parts.append(' | '.join(cells))
    return parts

def make_chunks(parts):
    chunks=[]
    current_title='ข้อมูลทั่วไป'
    buf=[]
    section_re=re.compile(r'^(\d+(?:\.\d+)?)\s*|^([๑-๙0-9]+\.)')
    def flush():
        nonlocal buf,current_title
        if buf:
            text='\n'.join(buf).strip()
            if len(text)>20:
                chunks.append({
                    'id': f'kkn_station_{len(chunks)+1:03d}',
                    'source': 'metadata_for_layout.docx / Design.pdf',
                    'station': 'ศูนย์อุตุนิยมวิทยาภาคตะวันออกเฉียงเหนือตอนบน',
                    'wmo': '48381',
                    'title': current_title,
                    'text': text
                })
            buf=[]
    for t in parts:
        is_heading=False
        if re.match(r'^\d+(\.\d+)?\s', t) or any(t.startswith(x) for x in ['ประวัติสถานี','2 รายละเอียด','3 ที่ตั้ง','4.รายละเอียด','5. สภาพ','6. การตรวจ','7. เครื่องมือ','8. ภูมิอากาศ','9. ค่ามาตรฐาน','10. สถิติ']):
            if len(t)<140:
                is_heading=True
        if is_heading:
            flush(); current_title=t; buf=[t]
        else:
            buf.append(t)
            if sum(len(x) for x in buf)>900:
                flush(); current_title=current_title+' (ต่อ)'
    flush()
    return chunks

parts=read_docx(SRC)
chunks=make_chunks(parts)
(BASE/'chunks.jsonl').write_text('\n'.join(json.dumps(c,ensure_ascii=False) for c in chunks),encoding='utf-8')
metadata={
 'name':'Khon Kaen Meteorological Station Vector Database',
 'language':'th',
 'station_th':'ศูนย์อุตุนิยมวิทยาภาคตะวันออกเฉียงเหนือตอนบน',
 'station_en':'UPPER NORTHEASTERN METEOROLOGICAL CENTER',
 'old_name':'สถานีตรวจอากาศขอนแก่น',
 'wmo':'48381',
 'local_rain_station':'381201',
 'location':'บริเวณท่าอากาศยานขอนแก่น ตำบลบ้านเป็ด อำเภอเมืองขอนแก่น จังหวัดขอนแก่น 40000',
 'coordinates':'ละติจูด 16°27\'39.7”N, ลองจิจูด 102°47\'22.6”E',
 'elevation_m':'186.97',
 'surface_observation_times':['01:00','04:00','07:00','10:00','13:00','16:00','19:00','22:00'],
 'upper_air_observation_times':['01:00','07:00','13:00','19:00'],
 'chunk_count':len(chunks),
 'embedding':'sklearn TfidfVectorizer analyzer=char_wb ngram_range=(2,5)',
 'created_from':['/mnt/data/metadata_for_layout.docx','/mnt/data/Design.pdf']
}
(BASE/'station_metadata.json').write_text(json.dumps(metadata,ensure_ascii=False,indent=2),encoding='utf-8')
texts=[c['title']+'\n'+c['text'] for c in chunks]
vectorizer=TfidfVectorizer(analyzer='char_wb', ngram_range=(2,5), max_features=50000)
X=vectorizer.fit_transform(texts)
nn=NearestNeighbors(n_neighbors=min(5,len(chunks)), metric='cosine').fit(X)
joblib.dump({'vectorizer':vectorizer,'matrix':X,'nn':nn,'chunks':chunks,'metadata':metadata}, BASE/'vector_db.joblib')
print('chunks',len(chunks))
